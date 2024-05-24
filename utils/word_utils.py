import os
import glob
import csv
from docx import Document
from .log_utils import log_error

def process_document(doc_path, log_path):
    try:
        doc = Document(doc_path)
        output = []

        for paragraph in doc.paragraphs:
            output.append(paragraph.text + "\n")

        for table in doc.tables:
            output.append("\n[Table Start]\n")
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.replace('\n', ' ').strip()
                    row_data.append(cell_text)
                output.append(" | ".join(row_data) + "\n")
            output.append("[Table End]\n")
        
        return ''.join(output)

    except Exception as e:
        log_error(log_path, doc_path, f"DOCX processing error: {e}")
        return ""

def save_text(output_text, output_path):
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(output_text)

def process_folder(folder_path, output_folder, log_path):
    for doc_path in glob.glob(os.path.join(folder_path, '*.docx')):
        output_file_name = os.path.basename(doc_path).replace('.docx', '.txt')
        output_path = os.path.join(output_folder, output_file_name)

        if os.path.exists(output_path):
            print(f"{output_path} already exists. Skipping.")
            continue
        
        structured_text = process_document(doc_path, log_path)
        
        if not structured_text.strip() or len(structured_text.strip()) < 20:
            log_error(log_path, doc_path, "Extracted text is empty or too short")
        
        save_text(structured_text, output_path)
        print(f"Saved processed text to {output_path}")

def extract_text_from_word_97(file_path):
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = True
    try:
        doc = word.Documents.Open(file_path)
        text = doc.Content.Text
        doc.Close()
        return text
    except Exception as e:
        raise Exception(f"Failed to process Word 97-2003 document: {e}")
    finally:
        word.Quit()

def process_word97_folder(input_folder, output_folder, log_path):
    for file_path in glob.glob(f"{input_folder}/*.doc"):
        output_file_name = os.path.basename(file_path).replace('.doc', '.txt')
        output_file_path = os.path.join(output_folder, output_file_name)

        if os.path.exists(output_file_path):
            print(f"{output_file_path} already exists. Skipping.")
            continue
        
        try:
            output_text = extract_text_from_word_97(file_path)
            with open(output_file_path, 'w', encoding='utf-8') as f:
                f.write(output_text)
            print(f"Processed and saved: {output_file_path}")

        except Exception as e:
            log_error(log_path, file_path, str(e))
            print(f"Error processing file '{file_path}': {e}")